import os
import re

import pdfplumber
from docx import Document

from app.service.ai_service import extract_resume_with_ai


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

def extract_text_from_pdf(file_path: str) -> str:

    if not os.path.exists(file_path):
        raise FileNotFoundError(
            "Resume file not found"
        )

    pages_text = []

    try:

        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                words = page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False
                )

                if not words:
                    continue

                # ------------------------------------------------
                # GROUP WORDS INTO VISUAL LINES
                # ------------------------------------------------

                lines = []

                for word in sorted(
                    words,
                    key=lambda w: (
                        w["top"],
                        w["x0"]
                    )
                ):

                    placed = False

                    for line in lines:

                        # Compare vertical position
                        if abs(
                            word["top"] - line["top"]
                        ) <= 4:

                            line["words"].append(
                                word
                            )

                            placed = True
                            break

                    if not placed:

                        lines.append({
                            "top": word["top"],
                            "words": [word]
                        })

                # ------------------------------------------------
                # SORT WORDS INSIDE EACH LINE
                # ------------------------------------------------

                for line in lines:

                    line["words"].sort(
                        key=lambda w: w["x0"]
                    )

                # ------------------------------------------------
                # SORT LINES TOP -> BOTTOM
                # ------------------------------------------------

                lines.sort(
                    key=lambda line: line["top"]
                )

                # ------------------------------------------------
                # BUILD TEXT
                # ------------------------------------------------

                page_lines = []

                for line in lines:

                    words_in_line = [
                        word["text"]
                        for word in line["words"]
                    ]

                    line_text = " ".join(
                        words_in_line
                    ).strip()

                    if line_text:

                        page_lines.append(
                            line_text
                        )

                pages_text.append(
                    "\n".join(page_lines)
                )

    except Exception as e:

        raise ValueError(
            f"Unable to read PDF: {str(e)}"
        )

    return "\n\n".join(
        pages_text
    ).strip()

# ============================================================
# DOCX TEXT EXTRACTION
# ============================================================

def extract_text_from_docx(file_path: str) -> str:

    if not os.path.exists(file_path):
        raise FileNotFoundError(
            "Resume file not found"
        )

    try:

        document = Document(
            file_path
        )

    except Exception as e:

        raise ValueError(
            f"Unable to read DOCX: {str(e)}"
        )

    text_parts = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            text_parts.append(
                text
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            row_parts = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_parts.append(
                        cell_text
                    )

            if row_parts:

                text_parts.append(
                    " | ".join(row_parts)
                )

    return "\n".join(
        text_parts
    ).strip()


# ============================================================
# GENERAL RESUME TEXT EXTRACTION
# ============================================================

def extract_resume_text(file_path: str) -> str:

    extension = os.path.splitext(
        file_path
    )[1].lower()

    if extension == ".pdf":

        return extract_text_from_pdf(
            file_path
        )

    elif extension == ".docx":

        return extract_text_from_docx(
            file_path
        )

    else:

        raise ValueError(
            "Only PDF and DOCX files are supported"
        )


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text: str) -> str:

    if not text:
        return ""

    # Normalize special spaces
    text = text.replace(
        "\xa0",
        " "
    )

    # Normalize line endings
    text = text.replace(
        "\r\n",
        "\n"
    )

    text = text.replace(
        "\r",
        "\n"
    )

    lines = []

    for line in text.split("\n"):

        # Keep useful spaces
        line = re.sub(
            r"[ \t]+",
            " ",
            line
        ).strip()

        if line:
            lines.append(
                line
            )

    # Rebuild text
    text = "\n".join(
        lines
    )

    # Prevent huge gaps
    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# ============================================================
# EMAIL FALLBACK
# ============================================================

def extract_email(text: str):

    if not text:
        return None

    pattern = (
        r"[A-Za-z0-9._%+\-]+"
        r"@"
        r"[A-Za-z0-9.\-]+"
        r"\."
        r"[A-Za-z]{2,}"
    )

    match = re.search(
        pattern,
        text
    )

    if match:

        return match.group(
            0
        ).lower().strip()

    return None


# ============================================================
# PHONE FALLBACK
# ============================================================

def extract_phone(text: str):

    if not text:
        return None

    # +91 9876543210
    match = re.search(
        r"(?:\+91|91)"
        r"[\s\-]*"
        r"([6-9]\d{4})"
        r"[\s\-]*"
        r"(\d{5})",
        text
    )

    if match:

        return (
            match.group(1)
            + match.group(2)
        )

    # 9876543210
    match = re.search(
        r"(?<!\d)"
        r"([6-9]\d{4})"
        r"[\s\-]*"
        r"(\d{5})"
        r"(?!\d)",
        text
    )

    if match:

        return (
            match.group(1)
            + match.group(2)
        )

    return None


# ============================================================
# DEFAULT DETAILS
# ============================================================

def empty_details():

    return {
        "name": None,
        "email": None,
        "phone": None,
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "languages": []
    }


# ============================================================
# CANDIDATE DETAILS
# ============================================================

def extract_candidate_details(text: str):

    cleaned_text = clean_text(
        text
    )

    if not cleaned_text:

        return empty_details()

    # ========================================================
    # LOCAL AI EXTRACTION
    # ========================================================
    #
    # IMPORTANT:
    # This calls our local ai_service.py.
    #
    # NO OpenAI
    # NO API KEY
    # NO INTERNET
    #
    # ========================================================

    details = extract_resume_with_ai(
        cleaned_text
    )

    # Safety check
    if not isinstance(
        details,
        dict
    ):

        details = empty_details()

    # ========================================================
    # Make sure all fields exist
    # ========================================================

    defaults = empty_details()

    for key, default_value in defaults.items():

        if key not in details:

            details[key] = default_value

    # ========================================================
    # Email fallback
    # ========================================================

    if not details.get(
        "email"
    ):

        details["email"] = extract_email(
            cleaned_text
        )

    # ========================================================
    # Phone fallback
    # ========================================================

    if not details.get(
        "phone"
    ):

        details["phone"] = extract_phone(
            cleaned_text
        )

    return details


# ============================================================
# COMPLETE RESUME PROCESSOR
# ============================================================

def process_resume(file_path: str):

    """
    Complete local resume processing.

    Returns:

    {
        "text": "...",
        "details": {
            "name": "...",
            "email": "...",
            "phone": "...",
            "skills": [],
            "education": [],
            "experience": [],
            "projects": [],
            "certifications": [],
            "languages": []
        }
    }
    """

    # --------------------------------------------------------
    # Extract PDF / DOCX
    # --------------------------------------------------------

    text = extract_resume_text(
        file_path
    )

    # --------------------------------------------------------
    # Clean extracted text
    # --------------------------------------------------------

    cleaned_text = clean_text(
        text
    )

    # --------------------------------------------------------
    # Local AI extraction
    # --------------------------------------------------------

    details = extract_candidate_details(
        cleaned_text
    )

    return {
        "text": cleaned_text,
        "details": details
    }